from __future__ import annotations

import json
import math
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


OUTPUT_DIR = Path(
    "/Users/abdul/.codex/visualizations/2026/07/18/"
    "019f76f7-ea1f-7493-9229-004c75d2bcf8/outputs/"
    "aurea-third-party-cost-report"
)
REPORT_PATH = Path(
    "/Users/abdul/Desktop/aurea-crm/docs/"
    "THIRD_PARTY_COST_MODEL_2026-07-18.md"
)
MODEL_PATH = OUTPUT_DIR / "model-results.json"
DOCX_PATH = OUTPUT_DIR / "Aurea_Third_Party_Technology_Cost_Report_2026-07-18.docx"
CHART_PATH = OUTPUT_DIR / "technology-cost-scaling.png"

NAVY = RGBColor(32, 55, 72)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
TEAL = RGBColor(15, 118, 110)
INK = RGBColor(32, 38, 46)
MUTED = RGBColor(93, 105, 116)
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
LIGHT_AMBER = "FFF4DF"
WHITE = RGBColor(255, 255, 255)
GRID = "D5DBE1"

PORTRAIT_WIDTH_DXA = 9360
LANDSCAPE_WIDTH_DXA = 14256
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_SIDE_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int, start: int, bottom: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], total_width: int) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        old = tbl_pr.find(qn(tag))
        if old is not None:
            tbl_pr.remove(old)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(total_width))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(
                cell,
                CELL_TOP_BOTTOM_DXA,
                CELL_SIDE_DXA,
                CELL_TOP_BOTTOM_DXA,
                CELL_SIDE_DXA,
            )


def set_table_borders(table, color: str = GRID, size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def set_font(run, name: str, size: float, color: RGBColor, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_font(run, "Calibri", 8.5, MUTED)


def set_header_footer(section, first_page: bool = False) -> None:
    section.different_first_page_header_footer = first_page
    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.clear()
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [4680, 4680], PORTRAIT_WIDTH_DXA)
    table._tbl.tblPr.append(OxmlElement("w:tblBorders"))
    left = table.cell(0, 0).paragraphs[0]
    right = table.cell(0, 1).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    left_run = left.add_run("AUREA | THIRD-PARTY COST REPORT")
    right_run = right.add_run("CONFIDENTIAL PLANNING")
    set_font(left_run, "Calibri", 8, MUTED, bold=True)
    set_font(right_run, "Calibri", 8, MUTED, bold=True)
    header_rule = header.add_paragraph()
    p_pr = header_rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    header_rule.paragraph_format.space_after = Pt(0)

    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    footer_table = footer.add_table(rows=1, cols=3, width=Inches(6.5))
    set_table_geometry(footer_table, [3120, 3120, 3120], PORTRAIT_WIDTH_DXA)
    footer_table._tbl.tblPr.append(OxmlElement("w:tblBorders"))
    values = ["18 July 2026", "Stakeholder planning model", ""]
    for index, value in enumerate(values):
        paragraph = footer_table.cell(0, index).paragraphs[0]
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.LEFT
            if index == 0
            else WD_ALIGN_PARAGRAPH.CENTER
            if index == 1
            else WD_ALIGN_PARAGRAPH.RIGHT
        )
        if index == 2:
            add_page_field(paragraph)
        else:
            run = paragraph.add_run(value)
            set_font(run, "Calibri", 8.5, MUTED)


def configure_section(section, landscape: bool = False) -> None:
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.75)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    set_header_footer(section, first_page=False)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    if "Table Citation" not in styles:
        style = styles.add_style("Table Citation", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["Table Citation"]
    style.font.name = "Calibri"
    style.font.size = Pt(8.5)
    style.font.color.rgb = MUTED
    style.paragraph_format.space_before = Pt(4)
    style.paragraph_format.space_after = Pt(4)

    if "Lead Callout" not in styles:
        style = styles.add_style("Lead Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["Lead Callout"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = NAVY
    style.font.bold = True
    style.paragraph_format.space_before = Pt(8)
    style.paragraph_format.space_after = Pt(10)
    style.paragraph_format.line_spacing = 1.15


def add_inline_markdown(paragraph, text: str, size: float | None = None) -> None:
    parts = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, "Calibri", size or 11, INK, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, "Aptos Mono", (size or 11) - 0.5, DARK_BLUE)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "EEF3F8")
            run._r.get_or_add_rPr().append(shading)
        else:
            run = paragraph.add_run(part)
            set_font(run, "Calibri", size or 11, INK)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [PORTRAIT_WIDTH_DXA], PORTRAIT_WIDTH_DXA)
    set_table_borders(table, color="D9A441", size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_AMBER)
    paragraph = cell.paragraphs[0]
    paragraph.style = doc.styles["Lead Callout"]
    add_inline_markdown(paragraph, text)
    paragraph.paragraph_format.keep_together = True
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def parse_markdown_table(lines: list[str], start: int):
    headers = [part.strip() for part in lines[start].strip().strip("|").split("|")]
    rows = []
    index = start + 2
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [
            part.strip().replace("\\|", "|")
            for part in lines[index].strip().strip("|").split("|")
        ]
        rows.append(cells)
        index += 1
    return headers, rows, index


def column_widths(headers: list[str], rows: list[list[str]], total: int) -> list[int]:
    weights = []
    for index, header in enumerate(headers):
        sample_lengths = [len(header)] + [
            len(row[index]) if index < len(row) else 0 for row in rows[:30]
        ]
        max_length = max(sample_lengths)
        numeric_ratio = sum(
            1
            for row in rows
            if index < len(row)
            and re.fullmatch(r"[\s£$€0-9.,%/()\-*]+", row[index] or "")
        ) / max(1, len(rows))
        if numeric_ratio > 0.7:
            weight = max(4.0, min(7.0, math.sqrt(max_length) * 1.6))
        else:
            weight = max(6.0, min(18.0, math.sqrt(max_length) * 2.4))
        weights.append(weight)
    raw = [int(total * weight / sum(weights)) for weight in weights]
    difference = total - sum(raw)
    raw[-1] += difference
    return raw


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    landscape: bool,
) -> None:
    total = LANDSCAPE_WIDTH_DXA if landscape else PORTRAIT_WIDTH_DXA
    widths = column_widths(headers, rows, total)
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths, total)
    set_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_repeat_table_header(table.rows[0])

    font_size = 7.3 if landscape else 8.4
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, LIGHT_GRAY)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(header)
        set_font(run, "Calibri", font_size, NAVY, bold=True)

    for row_data in rows:
        row = table.add_row()
        for index, cell_text in enumerate(row_data):
            if index >= len(row.cells):
                break
            cell = row.cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            numeric = bool(re.fullmatch(r"[\s£$€0-9.,%/()\-*]+", cell_text or ""))
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.RIGHT if numeric else WD_ALIGN_PARAGRAPH.LEFT
            )
            add_inline_markdown(paragraph, cell_text, size=font_size)

    following = doc.add_paragraph()
    following.paragraph_format.space_after = Pt(2)


def chart_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    font_name = "Arial Bold.ttf" if bold else "Arial.ttf"
    return ImageFont.truetype(
        f"/System/Library/Fonts/Supplemental/{font_name}",
        size,
    )


def money_tick(value: float) -> str:
    if value >= 1_000:
        return f"£{value / 1_000:,.0f}k"
    return f"£{value:,.0f}"


def build_chart() -> None:
    data = json.loads(MODEL_PATH.read_text())
    rows = [row for row in data["results"] if row["caseName"] == "Base"]
    locations = [row["locations"] for row in rows]
    labels = [f"{location:,}" for location in locations]
    revenue = [row["netRevenue"] for row in rows]
    core = [row["coreGbp"] for row in rows]
    comms = [row["communicationsGbp"] for row in rows]
    stripe = [row["stripeCollectionGbp"] for row in rows]
    current_cost = [row["currentCostGbp"] for row in rows]

    image = Image.new("RGB", (1872, 684), "white")
    draw = ImageDraw.Draw(image)
    title_font = chart_font(26, bold=True)
    label_font = chart_font(19)
    tick_font = chart_font(17)
    legend_font = chart_font(17)
    panels = [(70, 52, 894, 610), (978, 52, 1802, 610)]
    colors = {
        "navy": "#203748",
        "blue": "#2E74B5",
        "teal": "#0F766E",
        "amber": "#D9A441",
        "grid": "#D5DBE1",
        "muted": "#5D6974",
    }

    def draw_axes(
        panel: tuple[int, int, int, int],
        title: str,
        maximum: float,
    ) -> tuple[int, int, int, int]:
        x0, y0, x1, y1 = panel
        draw.text((x0, y0), title, fill=colors["navy"], font=title_font)
        plot = (x0 + 102, y0 + 116, x1 - 22, y1 - 55)
        px0, py0, px1, py1 = plot
        steps = 4
        for step in range(steps + 1):
            value = maximum * step / steps
            y = py1 - (py1 - py0) * step / steps
            draw.line((px0, y, px1, y), fill=colors["grid"], width=2)
            tick = money_tick(value)
            bbox = draw.textbbox((0, 0), tick, font=tick_font)
            draw.text(
                (px0 - (bbox[2] - bbox[0]) - 12, y - 10),
                tick,
                fill=colors["muted"],
                font=tick_font,
            )
        draw.line((px0, py0, px0, py1), fill=colors["muted"], width=2)
        draw.line((px0, py1, px1, py1), fill=colors["muted"], width=2)
        for item_index, label in enumerate(labels):
            x = px0 + (px1 - px0) * item_index / (len(labels) - 1)
            bbox = draw.textbbox((0, 0), label, font=tick_font)
            draw.text(
                (x - (bbox[2] - bbox[0]) / 2, py1 + 12),
                label,
                fill=colors["muted"],
                font=tick_font,
            )
        axis_label = "Studio locations"
        bbox = draw.textbbox((0, 0), axis_label, font=label_font)
        draw.text(
            ((px0 + px1 - (bbox[2] - bbox[0])) / 2, py1 + 38),
            axis_label,
            fill=colors["muted"],
            font=label_font,
        )
        return plot

    left_plot = draw_axes(
        panels[0],
        "Base-case revenue vs technology cost",
        math.ceil(max(revenue) / 100_000) * 100_000,
    )
    lx0, ly0, lx1, ly1 = left_plot
    left_max = math.ceil(max(revenue) / 100_000) * 100_000
    for values, color in (
        (revenue, colors["navy"]),
        (current_cost, colors["teal"]),
    ):
        points = []
        for item_index, value in enumerate(values):
            x = lx0 + (lx1 - lx0) * item_index / (len(values) - 1)
            y = ly1 - (ly1 - ly0) * value / left_max
            points.append((x, y))
        draw.line(points, fill=color, width=5, joint="curve")
        for x, y in points:
            draw.ellipse((x - 6, y - 6, x + 6, y + 6), fill=color)
    for legend_index, (name, color) in enumerate(
        (("Net SaaS revenue", colors["navy"]), ("Current tech cost", colors["teal"]))
    ):
        x = panels[0][0] + 106 + legend_index * 236
        y = panels[0][1] + 59
        draw.line((x, y + 9, x + 34, y + 9), fill=color, width=5)
        draw.text((x + 44, y), name, fill=colors["muted"], font=legend_font)

    right_max = math.ceil(max(current_cost) / 20_000) * 20_000
    right_plot = draw_axes(
        panels[1],
        "Current technology cost composition",
        right_max,
    )
    rx0, ry0, rx1, ry1 = right_plot
    slot = (rx1 - rx0) / len(labels)
    bar_width = min(48, slot * 0.6)
    series = (
        (core, colors["blue"]),
        (comms, colors["teal"]),
        (stripe, colors["amber"]),
    )
    for item_index in range(len(labels)):
        x = rx0 + slot * (item_index + 0.5)
        cumulative = 0.0
        for values, color in series:
            value = values[item_index]
            bottom_y = ry1 - (ry1 - ry0) * cumulative / right_max
            cumulative += value
            top_y = ry1 - (ry1 - ry0) * cumulative / right_max
            draw.rectangle(
                (x - bar_width / 2, top_y, x + bar_width / 2, bottom_y),
                fill=color,
            )
    legend_items = (
        ("Core", colors["blue"]),
        ("Managed communications", colors["teal"]),
        ("Stripe collection", colors["amber"]),
    )
    legend_x = panels[1][0] + 106
    for legend_index, (name, color) in enumerate(legend_items):
        x = legend_x + (0 if legend_index < 2 else 370)
        y = panels[1][1] + 58 + (legend_index if legend_index < 2 else 0) * 28
        draw.rectangle((x, y + 2, x + 18, y + 20), fill=color)
        draw.text((x + 28, y), name, fill=colors["muted"], font=legend_font)

    image.save(CHART_PATH, dpi=(180, 180))


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    configure_section(section, landscape=False)
    section.different_first_page_header_footer = True

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(116)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("STAKEHOLDER PLANNING REPORT")
    set_font(run, "Calibri", 10.5, BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run("Aurea Third-Party\nTechnology Cost Report")
    set_font(title_run, "Calibri", 30, NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    sub_run = subtitle.add_run(
        "Low, base and high scenarios across 1 to 1,000 Studio locations"
    )
    set_font(sub_run, "Calibri", 15, DARK_BLUE)

    offer = doc.add_paragraph()
    offer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    offer.paragraph_format.space_after = Pt(72)
    offer_run = offer.add_run("£500 per location per month, excluding VAT")
    set_font(offer_run, "Calibri", 11, BLUE, bold=True)

    metadata = doc.add_table(rows=3, cols=2)
    set_table_geometry(metadata, [2800, 6560], PORTRAIT_WIDTH_DXA)
    set_table_borders(metadata, color="E2E8F0", size="4")
    rows = [
        ("As of", "18 July 2026"),
        ("Status", "Planning model, not a vendor quote or audited forecast"),
        ("Scope", "Applications, infrastructure, communications, Stripe and customer offsets"),
    ]
    for row_index, (label, value) in enumerate(rows):
        set_cell_shading(metadata.cell(row_index, 0), LIGHT_GRAY)
        label_para = metadata.cell(row_index, 0).paragraphs[0]
        value_para = metadata.cell(row_index, 1).paragraphs[0]
        label_run = label_para.add_run(label)
        value_run = value_para.add_run(value)
        set_font(label_run, "Calibri", 9.5, NAVY, bold=True)
        set_font(value_run, "Calibri", 9.5, INK)

    footer_note = doc.add_paragraph()
    footer_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_note.paragraph_format.space_before = Pt(34)
    note_run = footer_note.add_run(
        "Confidential planning material | Refresh with actual invoices before contractual use"
    )
    set_font(note_run, "Calibri", 9, MUTED, italic=True)
    doc.add_page_break()


def add_contents(doc: Document, markdown: str) -> None:
    heading = doc.add_paragraph("Contents", style="Heading 1")
    set_keep_with_next(heading)
    for line in markdown.splitlines():
        match = re.match(r"^## (\d+\..+)$", line.strip())
        if not match:
            continue
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(match.group(1).split(". ", 1)[1])
        set_font(run, "Calibri", 10.5, INK)
    doc.add_page_break()


def render_markdown(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    body_started = False
    landscape_active = False
    chart_added = False

    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not body_started:
            if line.startswith("## 1."):
                body_started = True
            else:
                index += 1
                continue

        if not line:
            index += 1
            continue
        if line == "---":
            index += 1
            continue

        if line.startswith("## "):
            paragraph = doc.add_paragraph(line[3:], style="Heading 1")
            set_keep_with_next(paragraph)
            index += 1
            continue
        if line.startswith("### "):
            paragraph = doc.add_paragraph(line[4:], style="Heading 2")
            set_keep_with_next(paragraph)
            index += 1
            continue
        if line.startswith("#### "):
            paragraph = doc.add_paragraph(line[5:], style="Heading 3")
            set_keep_with_next(paragraph)
            index += 1
            continue

        if line.startswith("|") and index + 1 < len(lines):
            headers, rows, next_index = parse_markdown_table(lines, index)
            use_landscape = len(headers) >= 6
            moved_heading: tuple[str, str] | None = None
            if use_landscape and doc.paragraphs:
                previous = doc.paragraphs[-1]
                if previous.style.name.startswith("Heading"):
                    moved_heading = (previous.text, previous.style.name)
                    previous._element.getparent().remove(previous._element)
            if use_landscape:
                section = doc.add_section(WD_SECTION.NEW_PAGE)
                configure_section(section, landscape=True)
                landscape_active = True
                if moved_heading:
                    paragraph = doc.add_paragraph(
                        moved_heading[0],
                        style=moved_heading[1],
                    )
                    set_keep_with_next(paragraph)
            add_table(doc, headers, rows, landscape=use_landscape)
            if use_landscape:
                section = doc.add_section(WD_SECTION.NEW_PAGE)
                configure_section(section, landscape=False)
                landscape_active = False
            if not chart_added and "Net SaaS revenue" in headers and len(headers) >= 8:
                figure_para = doc.add_paragraph()
                figure_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                figure_para.add_run().add_picture(str(CHART_PATH), width=Inches(6.4))
                caption = doc.add_paragraph(
                    "Figure 1. Base-case revenue, technology cost and cost composition.",
                    style="Table Citation",
                )
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                chart_added = True
            index = next_index
            continue

        if re.match(r"^\d+\.\s", line):
            paragraph = doc.add_paragraph(style="List Number")
            add_inline_markdown(paragraph, re.sub(r"^\d+\.\s*", "", line))
            index += 1
            continue
        if line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, line[2:])
            index += 1
            continue

        paragraph_lines = [line]
        cursor = index + 1
        while cursor < len(lines):
            next_line = lines[cursor].strip()
            if (
                not next_line
                or next_line.startswith("#")
                or next_line.startswith("|")
                or next_line.startswith("- ")
                or re.match(r"^\d+\.\s", next_line)
                or next_line == "---"
            ):
                break
            paragraph_lines.append(next_line)
            cursor += 1
        text = " ".join(paragraph_lines)
        if text.startswith("The model supports four distinct conclusions:"):
            paragraph = doc.add_paragraph()
            add_inline_markdown(paragraph, text)
        elif text.startswith("The headline metric"):
            add_callout(doc, text)
        else:
            paragraph = doc.add_paragraph()
            add_inline_markdown(paragraph, text)
        index = cursor

    if landscape_active:
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        configure_section(section, landscape=False)


def build_document() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_chart()
    markdown = REPORT_PATH.read_text()
    doc = Document()
    configure_styles(doc)
    doc.core_properties.title = "Aurea Third-Party Technology Cost Report"
    doc.core_properties.subject = "Stakeholder technology cost planning"
    doc.core_properties.author = "Aurea"
    doc.core_properties.keywords = (
        "Aurea, third-party software, cost model, Stripe, Studios"
    )
    add_cover(doc)
    add_contents(doc, markdown)
    render_markdown(doc, markdown)
    doc.save(DOCX_PATH)
    print(f"DOCX={DOCX_PATH}")
    print(f"CHART={CHART_PATH}")


if __name__ == "__main__":
    build_document()
